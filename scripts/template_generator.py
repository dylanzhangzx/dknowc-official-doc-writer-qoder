#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公文红头文件生成器（代码化版）

设计逻辑：
1. 接收已生成的普通格式文档路径
2. 识别素材使用情况、参考资料和知识专库链接内容
3. 去除正文中的引用标记 [^1^] 等
4. 通过代码生成红头表格（开头）
5. 通过代码在正文末尾、素材使用情况或参考来源之前插入版记
6. 在版记后保留素材使用情况、参考资料和知识专库链接
7. 另存为红头版本
"""

import sys
import os
import re
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import argparse
import json

AI_DISCLAIMER_TEXT = "【AI生成提示】内容由AI生成，内容仅供参考。"
SKILL_ROOT = Path(__file__).resolve().parent.parent
OFFICIAL_DOCS_DIR = SKILL_ROOT / "official-docs"
INPUT_DIR = OFFICIAL_DOCS_DIR / "input"
OUTPUT_DIR = OFFICIAL_DOCS_DIR / "output"
PROFILE_PATH = SKILL_ROOT / "config" / "user_profile.json"
PLACEHOLDER_ORG = "XX单位"

UPWARD_TYPES = {"请示", "报告"}
DOWNWARD_TYPES = {"通知", "通报", "意见", "通告", "公告"}
LETTER_TYPES = {"函", "复函", "批复", "提醒函"}
MINUTES_TYPES = {"工作会议纪要", "局长办公会议纪要", "党组会议纪要", "会议纪要"}


def load_user_profile():
    """读取仅保存在本机的可选用户配置。"""
    if not PROFILE_PATH.exists():
        return {}
    try:
        with PROFILE_PATH.open("r", encoding="utf-8") as profile_file:
            profile = json.load(profile_file)
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"用户配置不可读取，请重新初始化: {exc}") from exc
    return profile if isinstance(profile, dict) else {}


def is_relative_to(path: Path, parent: Path) -> bool:
    """兼容旧 Python 版本的 Path.is_relative_to。"""
    try:
        path.relative_to(parent)
        return True
    except ValueError:
        return False


def resolve_input_docx(input_path) -> Path:
    """只允许读取 skill 工作目录内的 DOCX 文件。"""
    raw_path = Path(input_path).expanduser()
    if raw_path.is_absolute():
        resolved = raw_path.resolve()
    elif raw_path.parent == Path("."):
        resolved = (OUTPUT_DIR / raw_path.name).resolve()
    else:
        resolved = (SKILL_ROOT / raw_path).resolve()

    if resolved.suffix.lower() != ".docx":
        raise ValueError(f"只允许读取 .docx 文件: {input_path}")
    allowed_dirs = (INPUT_DIR.resolve(), OUTPUT_DIR.resolve())
    if not any(is_relative_to(resolved, allowed_dir) for allowed_dir in allowed_dirs):
        raise ValueError(f"输入文件必须位于 skill 工作目录内: {input_path}")
    return resolved


def resolve_output_docx(output_path) -> Path:
    """只允许将红头文档写入固定 Word 输出目录。"""
    raw_path = Path(output_path).expanduser()
    if raw_path.is_absolute():
        resolved = raw_path.resolve()
    elif raw_path.parent == Path("."):
        resolved = (OUTPUT_DIR / raw_path.name).resolve()
    else:
        resolved = (SKILL_ROOT / raw_path).resolve()

    if resolved.suffix.lower() != ".docx":
        resolved = resolved.with_suffix(".docx")
    if not is_relative_to(resolved, OUTPUT_DIR.resolve()):
        raise ValueError(f"输出文件必须位于 Word 输出目录内: {OUTPUT_DIR}")
    return resolved


def display_path(path: Path) -> str:
    """将 Skill 内文件路径转换为面向用户的相对路径。"""
    resolved = Path(path).expanduser().resolve()
    try:
        return str(resolved.relative_to(SKILL_ROOT))
    except ValueError:
        return str(resolved)


def remove_reference_markers_from_doc(doc):
    """去除文档中所有段落的引用标记 [^1^] [^2^] 等"""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and re.search(r'\[\^\d+\^\]', run.text):
                run.text = re.sub(r'\[\^\d+\^\]', '', run.text)
                count += 1
    return count


def remove_ai_disclaimer_from_doc(doc):
    """删除已有 AI 生成提示，避免影响红头文件版记排版。"""
    removed = 0
    for para in list(doc.paragraphs):
        if para.text.strip() == AI_DISCLAIMER_TEXT:
            para._p.getparent().remove(para._p)
            removed += 1
    return removed


def find_references_and_links_range(doc):
    """
    找到素材使用情况/参考资料和知识专库链接的段落范围，以及分页符位置
    
    Returns:
        tuple: (分页符索引, 素材或参考来源起始索引, 知识专库链接结束索引)
    """
    page_break_idx = None
    ref_start = None
    ref_end = None
    
    for i, para in enumerate(doc.paragraphs):
        # 检查是否有分页符
        for run in para.runs:
            if hasattr(run, '_element'):
                for br in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                    break_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    if break_type == 'page':
                        page_break_idx = i
        
        text = para.text.strip()
        
        # 检测素材使用情况或参考资料标题
        if ref_start is None:
            if (
                text in ["素材使用情况", "参考资料", "参考来源"]
                or text.startswith("【素材使用情况】")
                or text.startswith("【参考资料】")
                or text.startswith("【参考来源】")
            ):
                ref_start = i
        
        # 检测知识专库链接
        if ref_start is not None and ref_end is None:
            if text.startswith("http") or not text:
                continue
            elif text == "知识专库链接" or text.startswith("【知识专库链接】"):
                ref_end = i
    
    # 如果没找到知识专库链接，结束索引为文档末尾
    if ref_start is not None and ref_end is None:
        ref_end = len(doc.paragraphs)
    
    return page_break_idx, ref_start, ref_end


def is_minutes_type(doc_type):
    """判断是否为会议纪要类型"""
    return doc_type in MINUTES_TYPES


def get_layout_type(doc_type):
    """根据文种归类红头版式。"""
    if doc_type in UPWARD_TYPES:
        return "upward"
    if doc_type in LETTER_TYPES:
        return "letter"
    if doc_type in MINUTES_TYPES:
        return "minutes"
    return "downward"


def update_wps_textboxes(doc_elem, replacements):
    """更新文档中所有WPS文本框(wps:txbx)里的占位符"""
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_wps = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    
    count = 0
    txbxs = doc_elem.findall(f'.//{{{ns_wps}}}txbx')
    for txbx in txbxs:
        txbxContent = txbx.find(f'{{{ns_w}}}txbxContent')
        if txbxContent is None:
            continue
        for t in txbxContent.findall(f'.//{{{ns_w}}}t'):
            if t.text:
                for key, value in replacements.items():
                    placeholder = f"【{key}】"
                    if placeholder in t.text:
                        t.text = t.text.replace(placeholder, value)
                        count += 1
    return count


def update_minutes_footer_table(table, replacements):
    """更新纪要表尾单个表格中的占位符"""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if "【出席人员】" in text:
                        run.text = text.replace("【出席人员】", replacements.get("出席人员", ""))
                    elif "【列席人员】" in text:
                        run.text = text.replace("【列席人员】", replacements.get("列席人员", ""))
                    elif "【印发单位】" in text:
                        run.text = text.replace("【印发单位】", replacements.get("印发单位", ""))
                    elif "【印发日期】" in text:
                        run.text = text.replace("【印发日期】", replacements.get("印发日期", ""))
                    elif "【成文日期】" in text:
                        run.text = text.replace("【成文日期】", replacements.get("成文日期", ""))
                    elif "【抄送】" in text:
                        run.text = text.replace("【抄送】", replacements.get("抄送", ""))


def is_table_row_empty(row):
    """判断表格行是否所有单元格均为空（忽略空格）"""
    for cell in row.cells:
        for para in cell.paragraphs:
            if para.text.strip():
                return False
    return True


def remove_empty_rows(table):
    """删除表格中所有为空的行（倒序遍历避免索引偏移）"""
    for row in reversed(table.rows):
        if is_table_row_empty(row):
            row._tr.getparent().remove(row._tr)


def generate_red_header_document(doc_type, input_path, replacements, output_path):
    """生成红头文件"""
    input_path = resolve_input_docx(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"普通格式文档不存在: {input_path}")

    doc = Document(input_path)
    print(f"✓ 已加载普通格式文档: {input_path}")
    remove_ai_disclaimer_from_doc(doc)

    # 1. 找到分页符位置、参考资料和知识专库链接的范围
    page_break_idx, ref_start, ref_end = find_references_and_links_range(doc)
    if page_break_idx is not None:
        print(f"✓ 分页符位置: 段落 {page_break_idx}")
    print(f"✓ 素材/参考来源范围: 段落 {ref_start} 到 {ref_end}") if ref_start is not None else print("✓ 未找到素材/参考来源")

    # 2. 去除正文中的引用标记（参考资料之前的内容）
    if ref_start is not None:
        for para in doc.paragraphs[:ref_start]:
            for run in para.runs:
                if run.text and re.search(r'\[\^\d+\^\]', run.text):
                    run.text = re.sub(r'\[\^\d+\^\]', '', run.text)
    else:
        remove_reference_markers_from_doc(doc)
    print(f"✓ 已去除正文引用标记")

    # 3. 插入代码化红头和版记
    configure_document_page(doc)
    create_programmatic_red_header_table(doc, doc_type, replacements)
    print(f"✓ 已生成代码化红头: {get_layout_type(doc_type)}")

    insert_programmatic_footer_table(doc, replacements, page_break_idx, ref_start, doc_type)
    pos_desc = "分页符之前" if page_break_idx is not None else ("素材/参考来源之前" if ref_start is not None else "文档末尾")
    print(f"✓ 已生成代码化版记: {pos_desc}")

    # 5. 保存文件
    output_path = resolve_output_docx(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    return str(output_path)


def set_cell_text(cell, text, font_name='仿宋_GB2312', font_size=16, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格文字。"""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_cell_run(para, text, font_name='仿宋_GB2312', font_size=16, bold=False, color=None):
    """向单元格段落添加指定字体文字。"""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def set_signer_cell_text(cell, signer, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """按国标设置签发人字段：“签发人：”仿宋，姓名楷体。"""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1
    if signer:
        add_cell_run(para, "签发人：", font_name='仿宋_GB2312', font_size=16)
        add_cell_run(para, signer, font_name='楷体_GB2312', font_size=16)


def configure_document_page(doc):
    """设置红头文件页边距，采用党政机关公文常用 A4 版心。"""
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)


def set_table_borders(table, color='C00000', size='12', edges=None):
    """设置表格边框；edges 为 None 时设置全部边框。"""
    edges = ('top', 'left', 'bottom', 'right', 'insideH', 'insideV') if edges is None else edges
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        if edge in edges:
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), size)
        else:
            element.set(qn('w:val'), 'nil')
            element.set(qn('w:sz'), '0')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_table_border_edge(table, edge, val='single', size='8', color='000000'):
    """设置单条表格边框。"""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    element = borders.find(qn(f'w:{edge}'))
    if element is None:
        element = OxmlElement(f'w:{edge}')
        borders.append(element)
    element.set(qn('w:val'), val)
    element.set(qn('w:sz'), size)
    element.set(qn('w:space'), '0')
    element.set(qn('w:color'), color)


def set_standard_imprint_borders(table, has_middle_line=False):
    """设置国标版记横线：首末粗线，中间细线，无竖线。"""
    set_table_borders(table, color='000000', size='0', edges=())
    set_table_border_edge(table, 'top', size='8')
    set_table_border_edge(table, 'bottom', size='8')
    if has_middle_line:
        set_table_border_edge(table, 'insideH', size='6')


def set_table_width(table, width_cm):
    """设置表格宽度。"""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(width_cm * 567)))
    tblW.set(qn('w:type'), 'dxa')


def set_cell_width(cell, width_cm):
    """设置单元格宽度。"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')


def set_row_height(row, height_cm, exact=False):
    """设置表格行高。"""
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST


def merge_row(row):
    """合并一整行。"""
    merged = row.cells[0]
    for cell in row.cells[1:]:
        merged = merged.merge(cell)
    return row.cells[0]


def set_paragraph_border(para, edge='bottom', color='C00000', size='18'):
    """给段落设置边框，用于代码化红线。"""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    border = pBdr.find(qn(f'w:{edge}'))
    if border is None:
        border = OxmlElement(f'w:{edge}')
        pBdr.append(border)
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:space'), '6')
    border.set(qn('w:color'), color)


def set_paragraph_text(para, text, font_name='仿宋_GB2312', font_size=16, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置段落文字。"""
    para.clear()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def create_header_paragraph(text, font_name='仿宋_GB2312', font_size=16, color=None, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """创建红头段落 XML。"""
    temp_doc = Document()
    para = temp_doc.add_paragraph()
    set_paragraph_text(para, text, font_name=font_name, font_size=font_size, color=color, alignment=alignment)
    return para._p


def create_blank_paragraph(space_after=0):
    """创建空段落 XML。"""
    temp_doc = Document()
    para = temp_doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    return para._p


def create_red_line_paragraph():
    """创建红线段落 XML。"""
    temp_doc = Document()
    para = temp_doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(18)
    set_paragraph_border(para, color='C00000', size='18')
    return para._p


def create_programmatic_red_header_table(doc, doc_type, replacements):
    """生成代码化红头，按官方样本的表格结构和尺寸生成。"""
    layout = get_layout_type(doc_type)
    if layout != "minutes":
        return create_programmatic_official_header_table(doc, doc_type, replacements)
    return create_programmatic_minutes_header_table(doc, doc_type, replacements)


def create_programmatic_official_header_table(doc, doc_type, replacements):
    """普通公文红头：按官方格式样本代码化生成表格红头。"""
    org_name = replacements.get("发文机关", "")
    doc_number = replacements.get("发文字号", "")
    secrecy = add_fullwidth_space_if_needed(replacements.get("密级", ""))
    urgency = add_fullwidth_space_if_needed(replacements.get("紧急程度", ""))
    signer = replacements.get("签发人", "")
    layout = get_layout_type(doc_type)

    if layout == "letter":
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, 15.6)
        set_table_borders(table, color='C00000', size='18', edges=('bottom',))
        set_row_height(table.rows[0], 2.6)
        set_row_height(table.rows[1], 0.5, exact=True)
        set_row_height(table.rows[2], 1.0)
        set_row_height(table.rows[3], 1.0)
        merge_row(table.rows[0])
        merge_row(table.rows[1])
        top_parts = [part for part in (secrecy, urgency) if part]
        set_cell_text(table.rows[0].cells[0], org_name, font_name='仿宋_GB2312', font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[1].cells[0], "", font_name='仿宋_GB2312', font_size=8)
        set_cell_text(table.rows[2].cells[0], top_parts[0] if top_parts else "", font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[2].cells[1], doc_number, font_name='仿宋_GB2312', font_size=16, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.rows[3].cells[0], top_parts[1] if len(top_parts) > 1 else "", font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[3].cells[1], "", font_name='仿宋_GB2312', font_size=16)
    else:
        cols = 2 if layout == "upward" else 1
        table = doc.add_table(rows=3, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, 15.6)
        set_table_borders(table, color='C00000', size='18', edges=('bottom',))
        set_row_height(table.rows[0], 3.1, exact=True)
        set_row_height(table.rows[1], 1.55)
        set_row_height(table.rows[2], 2.9 if layout == "upward" else 2.57, exact=True)

        top_parts = [part for part in (secrecy, urgency) if part]
        top_cell = merge_row(table.rows[0]) if cols > 1 else table.rows[0].cells[0]
        title_cell = merge_row(table.rows[1]) if cols > 1 else table.rows[1].cells[0]
        set_cell_text(top_cell, "\n".join(top_parts), font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(title_cell, f"{org_name}文件" if org_name else "文件", font_name='方正小标宋简体', font_size=34, color=RGBColor(0xFF, 0x00, 0x00))
        if layout == "upward":
            set_cell_text(table.rows[2].cells[0], doc_number, font_name='仿宋_GB2312', font_size=16)
            set_signer_cell_text(table.rows[2].cells[1], signer)
        else:
            set_cell_text(table.rows[2].cells[0], doc_number, font_name='仿宋_GB2312', font_size=16)

    doc._body._body.insert(0, table._tbl)
    set_table_keep_together(table)
    return table


def create_programmatic_red_header_paragraphs(doc, doc_type, replacements):
    """普通公文红头：段落 + 红线，不使用红头表格。"""
    org_name = replacements.get("发文机关", "")
    doc_number = replacements.get("发文字号", "")
    secrecy = add_fullwidth_space_if_needed(replacements.get("密级", ""))
    urgency = add_fullwidth_space_if_needed(replacements.get("紧急程度", ""))
    signer = replacements.get("签发人", "")
    layout = get_layout_type(doc_type)

    elements = []
    top_parts = [part for part in (secrecy, urgency) if part]
    if top_parts:
        elements.append(create_header_paragraph('    '.join(top_parts), font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT))

    if layout == "letter":
        title_text = org_name
        title_size = 22
    else:
        title_text = f"{org_name}文件" if org_name else "文件"
        title_size = 24
    elements.append(create_header_paragraph(title_text, font_name='方正小标宋简体', font_size=title_size, color=RGBColor(0xC0, 0x00, 0x00)))

    if layout == "upward" and signer:
        number_text = f"{doc_number}        签发人：{signer}" if doc_number else f"签发人：{signer}"
    else:
        number_text = doc_number
    align = WD_ALIGN_PARAGRAPH.RIGHT if layout == "letter" else WD_ALIGN_PARAGRAPH.CENTER
    elements.append(create_header_paragraph(number_text, font_name='仿宋_GB2312', font_size=16, alignment=align))
    elements.append(create_red_line_paragraph())

    body = doc._body._body
    for element in reversed(elements):
        body.insert(0, element)
    return elements


def create_programmatic_minutes_header_table(doc, doc_type, replacements):
    """会议纪要红头表格。"""
    org_name = replacements.get("发文机关", "")
    doc_number = replacements.get("发文字号", "")
    secrecy = add_fullwidth_space_if_needed(replacements.get("密级", ""))
    urgency = add_fullwidth_space_if_needed(replacements.get("紧急程度", ""))
    signer = replacements.get("签发人", "")
    meeting_number = replacements.get("纪要编号", "")

    table = doc.add_table(rows=4, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 15.6)
    set_table_borders(table, color='C00000', size='18', edges=('bottom',))
    set_row_height(table.rows[0], 0.45)
    set_row_height(table.rows[1], 1.15)
    for row in table.rows[2:]:
        set_row_height(row, 0.55)

    top_parts = [part for part in (secrecy, urgency) if part]
    set_cell_text(table.rows[0].cells[0], '    '.join(top_parts), font_name='黑体', font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    title_text = build_minutes_header_title(org_name, doc_type)
    set_cell_text(table.rows[1].cells[0], title_text, font_name='方正小标宋简体', font_size=22, color=RGBColor(0xC0, 0x00, 0x00))

    set_cell_text(table.rows[2].cells[0], meeting_number or doc_number, font_name='仿宋_GB2312', font_size=16)
    set_cell_text(table.rows[3].cells[0], replacements.get("印发单位", "") or f"{org_name}办公室", font_name='仿宋_GB2312', font_size=15, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    doc._body._body.insert(0, table._tbl)
    set_table_keep_together(table)
    return table


def build_minutes_header_title(org_name, doc_type):
    """生成会议纪要红头标题。"""
    if doc_type == "党组会议纪要":
        return f"中共{org_name}\n党组会议纪要" if org_name and not org_name.startswith("中共") else f"{org_name}\n党组会议纪要"
    if doc_type == "局长办公会议纪要":
        return f"{org_name}\n局长办公会议纪要" if org_name else "局长办公会议纪要"
    if doc_type == "工作会议纪要":
        return f"{org_name}\n工作会议纪要" if org_name else "工作会议纪要"
    return f"{org_name}\n会议纪要" if org_name else "会议纪要"


def add_spacing_after_programmatic_header(doc):
    """给程序化红头和正文之间补一个空段落。"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(12)
    body = doc._body._body
    # 红头表格已在最前方，空段落移动到红头后、正文前。
    body.insert(1, para._p)


def paragraph_has_text(para):
    """判断段落是否含有可见文本。"""
    return bool(para.text.strip())


def remove_page_breaks_from_paragraph(para):
    """移除段落中的分页符，用于把日期段落和素材附录分页拆开。"""
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    removed = False
    for br in list(para._element.findall(f'.//{{{ns_w}}}br')):
        if br.get(f'{{{ns_w}}}type') == 'page':
            parent = br.getparent()
            if parent is not None:
                parent.remove(br)
                removed = True
    return removed


def create_page_break_paragraph():
    """创建仅包含分页符的段落 XML。"""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def normalize_cc_text(cc_text):
    """规范版记抄送行，末尾补句号。"""
    text = (cc_text or "").strip()
    if not text:
        return ""
    if text[-1] not in "。.!！?？":
        text += "。"
    return text


def insert_xml_before_appendix(doc, xml_elements, page_break_idx=None, ref_start=None):
    """
    将版记 XML 插入正文和素材附录之间。
    如果分页符和日期同段，先移除原分页符，再在版记后补一个分页符，避免日期跑到版记后面。
    """
    if not isinstance(xml_elements, (list, tuple)):
        xml_elements = [xml_elements]

    if page_break_idx is not None:
        para = doc.paragraphs[page_break_idx]
        if paragraph_has_text(para):
            remove_page_breaks_from_paragraph(para)
            anchor = para._element
            for element in xml_elements:
                anchor.addnext(element)
                anchor = element
            anchor.addnext(create_page_break_paragraph())
        else:
            anchor = para._element
            for element in xml_elements:
                anchor.addprevious(element)
        return

    if ref_start is not None:
        anchor = doc.paragraphs[ref_start]._element
        for element in xml_elements:
            anchor.addprevious(element)
        return

    for element in xml_elements:
        doc._body._body.append(element)


def insert_programmatic_footer_table(doc, replacements, page_break_idx=None, ref_start=None, doc_type=None):
    """生成国标版记，并插入到素材页之前或正文末尾。"""
    if is_minutes_type(doc_type):
        return insert_programmatic_minutes_footer(doc, replacements, page_break_idx, ref_start)

    org_name = replacements.get("发文机关", "")
    date_str = replacements.get("印发日期") or replacements.get("成文日期", "")
    cc_list = replacements.get("抄送", "")
    layout = get_layout_type(doc_type)

    if layout == "letter":
        cc_text = normalize_cc_text(cc_list)
        if not cc_text:
            return None
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, 15.6)
        set_table_borders(table, color='000000', size='0', edges=())
        set_cell_text(table.rows[0].cells[0], f"抄送：{cc_text}", font_size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_keep_together(table)
        insert_xml_before_appendix(doc, table._tbl, page_break_idx, ref_start)
        return table

    cc_text = normalize_cc_text(cc_list)
    rows = 2 if cc_text else 1
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 15.6)
    set_standard_imprint_borders(table, has_middle_line=bool(cc_text))

    issue_row_idx = 0
    if cc_text:
        cc_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
        set_cell_text(cc_cell, f"抄送：{cc_text}", font_size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        issue_row_idx = 1

    issue_row = table.rows[issue_row_idx]
    set_cell_width(issue_row.cells[0], 9.0)
    set_cell_width(issue_row.cells[1], 6.6)
    print_unit = replacements.get("印发单位", "") or (f"{org_name}办公室" if org_name else "")
    set_cell_text(issue_row.cells[0], print_unit, font_size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(issue_row.cells[1], f"{date_str}印发" if date_str else "", font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    set_table_keep_together(table)
    insert_xml_before_appendix(doc, table._tbl, page_break_idx, ref_start)
    return table


def insert_programmatic_minutes_footer(doc, replacements, page_break_idx=None, ref_start=None):
    """生成会议纪要出席/列席信息和版记。"""
    attendees = replacements.get("出席人员", "")
    non_voting = replacements.get("列席人员", "")
    org_name = replacements.get("发文机关", "")
    print_unit = replacements.get("印发单位", "") or f"{org_name}办公室"
    print_date = replacements.get("印发日期", "") or replacements.get("成文日期", "")
    cc_list = replacements.get("抄送", "")

    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(info_table, 15.6)
    set_table_borders(info_table, color='000000', size='8', edges=('top', 'insideH', 'bottom'))
    set_cell_text(info_table.rows[0].cells[0], "出席：", font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(info_table.rows[0].cells[1], attendees, font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(info_table.rows[1].cells[0], "列席：", font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(info_table.rows[1].cells[1], non_voting, font_size=16, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    cc_text = normalize_cc_text(cc_list)
    issue_rows = 2 if cc_text else 1
    issue_table = doc.add_table(rows=issue_rows, cols=2)
    issue_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(issue_table, 15.6)
    set_standard_imprint_borders(issue_table, has_middle_line=bool(cc_text))
    issue_row_idx = 0
    if cc_text:
        cc_cell = issue_table.rows[0].cells[0].merge(issue_table.rows[0].cells[1])
        set_cell_text(cc_cell, f"抄送：{cc_text}", font_size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        issue_row_idx = 1
    set_cell_width(issue_table.rows[issue_row_idx].cells[0], 9.0)
    set_cell_width(issue_table.rows[issue_row_idx].cells[1], 6.6)
    set_cell_text(issue_table.rows[issue_row_idx].cells[0], print_unit, font_size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(issue_table.rows[issue_row_idx].cells[1], f"{print_date}印发" if print_date else "", font_size=14, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    for table in (info_table, issue_table):
        set_table_keep_together(table)

    insert_xml_before_appendix(doc, [info_table._tbl, issue_table._tbl], page_break_idx, ref_start)
    return info_table


def add_fullwidth_space_if_needed(text):
    """为两字密级/紧急程度自动添加全角空格"""
    if text and len(text) == 2:
        if all('\u4e00' <= c <= '\u9fff' for c in text):
            return text[0] + '\u3000' + text[1]
    return text


def update_red_header_table(table, replacements):
    """更新红头表格中的占位符"""
    org_name = replacements.get("发文机关", "")
    doc_number = replacements.get("发文字号", "")
    secrecy = add_fullwidth_space_if_needed(replacements.get("密级", ""))
    urgency = add_fullwidth_space_if_needed(replacements.get("紧急程度", ""))
    signer = replacements.get("签发人", "")
    meeting_number = replacements.get("纪要编号", "")
    print_unit = replacements.get("印发单位", "")
    print_date = replacements.get("印发日期", "")

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if "【发文机关】" in text:
                        run.text = text.replace("【发文机关】", org_name)
                    if "【发文字号】" in text:
                        run.text = text.replace("【发文字号】", doc_number)
                    if "【密级】" in text:
                        run.text = text.replace("【密级】", secrecy if secrecy else "")
                    if "【紧急程度】" in text:
                        run.text = text.replace("【紧急程度】", urgency if urgency else "")
                    if "【签发人】" in text:
                        run.text = text.replace("【签发人】", signer if signer else "")
                    if "【纪要编号】" in text:
                        run.text = text.replace("【纪要编号】", meeting_number)
                    if "【印发单位】" in text:
                        run.text = text.replace("【印发单位】", print_unit)
                    if "【成文日期】" in text and "印发" not in text:
                        run.text = text.replace("【成文日期】", print_date)


def update_vml_shapes(doc, replacements):
    """更新VML图形中的占位符"""
    org_name = replacements.get("发文机关", "")
    doc_number = replacements.get("发文字号", "")
    secrecy = replacements.get("密级", "")
    urgency = replacements.get("紧急程度", "")
    meeting_number = replacements.get("纪要编号", "")
    print_unit = replacements.get("印发单位", "")
    print_date = replacements.get("印发日期", "")

    body = doc._body._body
    text_nodes = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')

    for t in text_nodes:
        if t.text:
            if "【发文机关】" in t.text:
                t.text = t.text.replace("【发文机关】", org_name)
            if "【发文字号】" in t.text:
                t.text = t.text.replace("【发文字号】", doc_number)
            if "【密级】" in t.text:
                t.text = t.text.replace("【密级】", secrecy if secrecy else "")
            if "【紧急程度】" in t.text:
                t.text = t.text.replace("【紧急程度】", urgency if urgency else "")
            if "【纪要编号】" in t.text:
                t.text = t.text.replace("【纪要编号】", meeting_number)
            if "【印发单位】" in t.text:
                t.text = t.text.replace("【印发单位】", print_unit)
            if "【成文日期】" in t.text and "印发" not in t.text:
                t.text = t.text.replace("【成文日期】", print_date)


def set_table_keep_together(table):
    """设置表格不允许跨页拆分"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    keep_together = OxmlElement('w:keep-together')
    keep_together.set(qn('w:val'), '1')
    tblPr.append(keep_together)


def update_footer_table(table, replacements):
    """更新结尾表格中的占位符"""
    org_name = replacements.get("发文机关", "")
    date_str = replacements.get("成文日期", "")
    cc_list = replacements.get("抄送", "")

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    text = run.text
                    if "【发文机关】" in text:
                        run.text = text.replace("【发文机关】", org_name)
                    if "【成文日期】" in text:
                        run.text = text.replace("【成文日期】", date_str)
                    if "【印发日期】" in text:
                        run.text = text.replace("【印发日期】", date_str)
                    if "【公开方式】" in text:
                        run.text = text.replace("【公开方式】", "主动公开")
                    if "【抄送】" in text:
                        run.text = text.replace("【抄送】", cc_list)


def main():
    profile = load_user_profile()
    year = datetime.now().year
    doc_prefix = profile.get('doc_prefix') or "XX"
    default_doc_number = f"{doc_prefix}〔{year}〕XX号"
    parser = argparse.ArgumentParser(description='红头文件生成器 v4.4')
    parser.add_argument('type', help='文种')
    parser.add_argument('--input', required=True, help='普通格式文档路径')
    parser.add_argument('--org', default=profile.get('organization') or PLACEHOLDER_ORG, help='发文机关')
    parser.add_argument('--doc-number', default=default_doc_number, help='发文字号')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--signer', default='', help='签发人')
    # 会议纪要专用参数
    parser.add_argument('--meeting-number', default='', help='纪要编号（如〔2026〕第10号）')
    parser.add_argument('--attendees', default='', help='出席人员')
    parser.add_argument('--non-voting', default='', help='列席人员')
    parser.add_argument('--print-unit', default=profile.get('print_unit') or '', help='印发单位')
    parser.add_argument('--print-date', default='', help='印发日期（如2026年3月19日）')
    parser.add_argument('--cc', default='', help='抄送单位')

    args = parser.parse_args()

    input_path = resolve_input_docx(args.input)
    if args.output:
        output_path = args.output
    else:
        output_path = str(OUTPUT_DIR / f"{input_path.stem}_红头{input_path.suffix}")

    now = datetime.now()
    today = f"{now.year}年{now.month}月{now.day}日"

    replacements = {
        "发文机关": args.org,
        "发文字号": args.doc_number,
        "成文日期": today,
        "签发人": args.signer,
        # 会议纪要专用
        "纪要编号": args.meeting_number,
        "出席人员": args.attendees,
        "列席人员": args.non_voting,
        "印发单位": args.print_unit if args.print_unit else f"{args.org}办公室",
        "印发日期": args.print_date if args.print_date else today,
        "抄送": args.cc,
    }

    output = generate_red_header_document(args.type, input_path, replacements, output_path)
    print(f"✅ 红头文件已生成: {display_path(Path(output))}")


if __name__ == "__main__":
    main()
