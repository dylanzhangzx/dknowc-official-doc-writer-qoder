#!/usr/bin/env python3
"""
公文排版脚本 v2.0.0
严格按照《XX单位公文格式样本（试行）》进行自动排版

【v2.0.0 更新说明】
- 删除所有红头文件相关代码（红头+尾表由 template_generator.py 处理）
- 只保留普通格式排版核心逻辑
- 架构：template_generator.py 调用本文件的函数进行正文排版
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import docx
import docx.opc.constants
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url, text, font_name=None, font_size=None):
    """在段落中添加可点击的超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if font_name or font_size:
        rFonts = OxmlElement('w:rFonts')
        if font_name:
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(szCs)

    # 超链接颜色（蓝色）
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # 下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def sanitize_url(raw_url):
    """清理 URL 后缀说明，避免把括号内说明写进超链接。"""
    url = raw_url.strip().rstrip('，。、；;,.')
    url = re.split(r'[（(]', url, maxsplit=1)[0]
    return url.strip().rstrip('，。、；;,.')


def hyperlink_display_text(label=None, fallback_index=None):
    """生成知识专库链接的可见文本。"""
    clean_label = normalize_kb_label(label)
    if not clean_label:
        clean_label = f"知识专库链接{fallback_index}" if fallback_index else "知识专库链接"
    return f"{clean_label}（点击打开）"


def normalize_kb_label(label):
    """清洗知识专库链接标签，去掉编号、Markdown 强调和已写入的点击提示。"""
    clean_label = (label or "").strip()
    clean_label = re.sub(r'^\s*\d+[\.、]\s*', '', clean_label)
    clean_label = re.sub(r'^\s*[-*]\s*', '', clean_label)
    clean_label = re.sub(r'[*_`]+', '', clean_label)
    clean_label = re.sub(r'[（(]\s*点击打开\s*[）)]', '', clean_label)
    clean_label = clean_label.replace('点击打开', '')
    clean_label = re.sub(r'\s+', ' ', clean_label)
    return clean_label.strip().rstrip('：:')


def is_internal_search_endpoint(url):
    """识别深知搜索接口地址，避免误当成知识专库链接输出。"""
    return "open.dknowc.cn/dependable/search" in url


def is_generic_kb_label(text):
    """识别不能作为知识专库链接显示名的泛化说明。"""
    clean_text = (text or "").strip().rstrip('：:')
    if not clean_text:
        return True
    generic_patterns = (
        r'^以下.*知识专库链接.*',
        r'^本次.*知识专库.*',
        r'^知识专库链接$',
        r'^链接$',
        r'^来源链接$',
        r'^可点击链接$',
    )
    return any(re.match(pattern, clean_text) for pattern in generic_patterns)


def is_useless_kb_label_line(text):
    """识别知识专库链接区中可直接丢弃的无意义标签行。"""
    clean_text = (text or "").strip().rstrip('：:')
    return clean_text in {"知识专库链接", "链接", "来源链接", "可点击链接"}

DEFAULT_OUTPUT_DIR = "official-docs/output"
CONFIG_PATH = os.path.join(os.path.dirname(__file__), '..', 'config', 'format.json')
SKILL_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OFFICIAL_DOCS_DIR = os.path.join(SKILL_ROOT, "official-docs")
INPUT_DIR = os.path.join(OFFICIAL_DOCS_DIR, "input")
SEARCH_RESULTS_DIR = os.path.join(OFFICIAL_DOCS_DIR, "search-results")
AI_DISCLAIMER_TEXT = "【AI生成提示】内容由AI生成，内容仅供参考。"


def load_format_config():
    """加载格式配置文件"""
    try:
        with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"警告: 配置文件不存在 {CONFIG_PATH}，使用默认值")
        return None
    except json.JSONDecodeError as e:
        print(f"警告: 配置文件格式错误 {CONFIG_PATH}: {e}，使用默认值")
        return None


# 加载配置（全局变量）
FORMAT_CONFIG = load_format_config()


def get_skill_root():
    """Skill 安装根目录（含 scripts/、config/ 等），与调用时的 cwd 无关。"""
    return SKILL_ROOT


def get_configured_output_dir():
    """获取 Word 输出目录：相对路径基于 Skill 安装目录解析，不依赖 cwd。"""
    configured_dir = None
    if FORMAT_CONFIG:
        output_config = FORMAT_CONFIG.get('output', {})
        if isinstance(output_config, dict):
            configured_dir = output_config.get('dir') or output_config.get('output_dir')

    output_dir = os.path.expanduser(configured_dir or DEFAULT_OUTPUT_DIR)
    if not os.path.isabs(output_dir):
        output_dir = os.path.join(get_skill_root(), output_dir)
    output_dir = os.path.abspath(output_dir)
    if not is_path_within(output_dir, os.path.abspath(os.path.join(get_skill_root(), DEFAULT_OUTPUT_DIR))):
        raise ValueError(f"输出目录必须位于 {os.path.join(get_skill_root(), DEFAULT_OUTPUT_DIR)} 内")
    return output_dir


def is_path_within(path, parent):
    """判断 path 是否位于 parent 目录内。"""
    try:
        return os.path.commonpath([os.path.abspath(path), os.path.abspath(parent)]) == os.path.abspath(parent)
    except ValueError:
        return False


def resolve_input_text_path(input_path):
    """只允许读取 skill 工作目录内的正文文本文件。"""
    raw_path = os.path.expanduser(str(input_path).strip())
    if os.path.isabs(raw_path):
        resolved = os.path.abspath(raw_path)
    elif raw_path == os.path.basename(raw_path):
        resolved = os.path.abspath(os.path.join(INPUT_DIR, raw_path))
    else:
        resolved = os.path.abspath(os.path.join(get_skill_root(), raw_path))

    allowed_dirs = [
        os.path.abspath(INPUT_DIR),
        os.path.abspath(get_configured_output_dir()),
        os.path.abspath(SEARCH_RESULTS_DIR),
    ]
    if not any(is_path_within(resolved, allowed_dir) for allowed_dir in allowed_dirs):
        raise ValueError(f"输入文件必须位于 skill 工作目录内: {input_path}")
    if os.path.splitext(resolved)[1].lower() not in {".txt", ".md"}:
        raise ValueError(f"只允许读取 .txt 或 .md 正文文件: {input_path}")
    return resolved


def resolve_output_path(output_path):
    """解析输出路径：只允许写入固定 Word 输出目录。"""
    if not output_path:
        return None
    output_path = os.path.expanduser(output_path.strip())
    output_dir = get_configured_output_dir()
    if os.path.isabs(output_path):
        resolved = os.path.abspath(output_path)
    elif output_path == os.path.basename(output_path):
        resolved = os.path.abspath(os.path.join(output_dir, output_path))
    else:
        resolved = os.path.abspath(os.path.join(get_skill_root(), output_path))
    if not is_path_within(resolved, output_dir):
        raise ValueError(f"输出文件必须位于 Word 输出目录内: {output_dir}")
    return resolved


def display_path(path):
    """将 Skill 内文件路径转换为面向用户的相对路径。"""
    resolved = os.path.abspath(os.path.expanduser(str(path)))
    try:
        return os.path.relpath(resolved, get_skill_root())
    except ValueError:
        return resolved


def get_font_config(element_name):
    """从配置获取字体设置，返回 (font_name, font_size, bold)"""
    if not FORMAT_CONFIG or element_name not in FORMAT_CONFIG:
        # 默认值
        defaults = {
            'title': ('方正小标宋简体', 22, False),
            'body': ('仿宋_GB2312', 16, False),
            'heading1': ('黑体', 16, False),
            'heading2': ('楷体_GB2312', 16, False),
            'heading3': ('仿宋_GB2312', 16, False),
        }
        return defaults.get(element_name, ('仿宋_GB2312', 16, False))
    
    config = FORMAT_CONFIG[element_name]
    return (
        config.get('font', '仿宋_GB2312'),
        config.get('size_pt', 16),
        config.get('bold', False)
    )


def get_page_margin():
    """从配置获取页边距，返回 (top_mm, bottom_mm, left_mm, right_mm)"""
    if FORMAT_CONFIG and 'page' in FORMAT_CONFIG:
        page = FORMAT_CONFIG['page']
        # 支持四边分别设置（国标 GB/T 9704）
        if 'margin_top_mm' in page:
            return (
                page.get('margin_top_mm', 37),
                page.get('margin_bottom_mm', 35),
                page.get('margin_left_mm', 28),
                page.get('margin_right_mm', 26),
            )
        # 兼容旧配置：统一边距
        unified = page.get('margin_mm', 25)
        return (unified, unified, unified, unified)
    return (37, 35, 28, 26)


def get_page_size():
    """从配置获取纸张尺寸，默认 A4。"""
    if FORMAT_CONFIG and 'page' in FORMAT_CONFIG:
        page = FORMAT_CONFIG['page']
        return (
            page.get('width_mm', 210),
            page.get('height_mm', 297),
        )
    return (210, 297)


def get_latin_font_config():
    """获取全文数字和字母字体。"""
    if FORMAT_CONFIG:
        body = FORMAT_CONFIG.get('body', {})
        if body.get('latin_font'):
            return body['latin_font']
    return 'Times New Roman'

MAX_INPUT_LENGTH = 50000
MAX_INPUT_LINES = 2000
MAX_OUTPUT_LINES = 3000


def get_next_version(output_path: str) -> str:
    """
    获取下一个版本号的输出路径

    Args:
        output_path: 原始输出路径

    Returns:
        带版本号的输出路径（如果文件已存在）
    """
    output_path = os.path.expanduser(output_path)
    output_dir = Path(output_path).parent
    filename = Path(output_path).stem  # 文件名（无扩展名）
    ext = Path(output_path).suffix     # 扩展名

    # 如果文件不存在，直接返回原路径
    if not Path(output_path).exists():
        return output_path

    # 文件已存在，需要添加版本号
    # 查找所有同名文件（包含版本号）
    pattern = f"{filename}_v*.docx"
    existing_files = list(output_dir.glob(pattern))

    # 如果没有版本号文件，从v1开始
    if not existing_files:
        return str(output_dir / f"{filename}_v1{ext}")

    # 提取所有版本号
    versions = []
    for file in existing_files:
        # 提取版本号（如"关于XXX的通知_v1.docx" -> 1）
        try:
            version_str = file.stem.split('_v')[-1]
            version_num = int(version_str)
            versions.append(version_num)
        except:
            continue

    # 计算下一个版本号
    next_version = max(versions) + 1 if versions else 1

    return str(output_dir / f"{filename}_v{next_version}{ext}")


def sanitize_filename_component(text: str, max_length: int = 80) -> str:
    """将公文标题转换为可用文件名。"""
    value = re.sub(r'\s+', '', text.strip())
    value = re.sub(r'[\\/:*?"<>|]', '', value)
    value = value.strip('. ')
    if not value:
        return ""
    return value[:max_length]


try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 缺少 python-docx 库，请安装: pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_name, font_size, bold=False, color=None, latin_font=None, latin=False):
    """设置字体"""
    latin_font = latin_font or get_latin_font_config()
    display_font = latin_font if latin else font_name
    run.font.name = display_font
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), display_font)
    rFonts.set(qn('w:hAnsi'), display_font)
    rFonts.set(qn('w:cs'), display_font)


def add_formatted_text(paragraph, text, font_name, font_size, bold=False, color=None):
    """添加文本并确保数字、字母使用 Times New Roman。"""
    latin_font = get_latin_font_config()
    for segment in re.findall(r'[A-Za-z0-9]+|[^A-Za-z0-9]+', text):
        run = paragraph.add_run(segment)
        set_run_font(
            run,
            font_name,
            font_size,
            bold=bold,
            color=color,
            latin_font=latin_font,
            latin=bool(re.fullmatch(r'[A-Za-z0-9]+', segment)),
        )
    return paragraph


def set_paragraph_format(para, first_line_indent=False, indent_chars=None, alignment=None, line_spacing=None):
    """设置段落格式
    
    Args:
        para: 段落对象
        first_line_indent: 是否首行缩进（已废弃，请使用indent_chars）
        indent_chars: 缩进字符数（如2表示2字符）
        alignment: 对齐方式
        line_spacing: 行距（磅）
    """
    if indent_chars:
        # 同时写入常规缩进，兼容只读取 first_line_indent 的校验器。
        para.paragraph_format.first_line_indent = Cm(0.37 * indent_chars)
        # 使用Word的字符单位缩进（更准确）
        # 通过底层XML设置 w:firstLineChars 属性
        # 值的单位是1/100字符，所以2字符 = 200
        pPr = para._p.get_or_add_pPr()
        ind = pPr.ind
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(int(indent_chars * 100)))
    elif first_line_indent:
        # 兼容旧配置
        indent_cm = FORMAT_CONFIG['body']['first_line_indent_cm'] if FORMAT_CONFIG else 0.74
        para.paragraph_format.first_line_indent = Cm(indent_cm)
    if alignment is not None:
        para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def set_outline_level(para, level):
    """设置段落的大纲级别
    
    Args:
        para: 段落对象
        level: 大纲级别（0-9，0为最高级）
               0 = 1级（对应 # 标题）
               1 = 2级（对应 ## 一、XXX）
               2 = 3级（对应 ### （一）XXX）
               3 = 4级（对应 #### 1.XXX）
               9 = 正文文本（默认）
    """
    pPr = para._p.get_or_add_pPr()
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level))
    pPr.append(outlineLvl)


def add_page_field(run):
    """添加 Word 页码域。"""
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def add_standard_page_number(footer, alignment):
    """按 GB/T 9704 页码样式添加页脚：数字左右各一字线。"""
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = alignment
    left_dash = para.add_run("— ")
    set_run_font(left_dash, '宋体', 14)
    run = para.add_run()
    set_run_font(run, '宋体', 14)
    add_page_field(run)
    right_dash = para.add_run(" —")
    set_run_font(right_dash, '宋体', 14)


def add_gbt_page_numbers(doc, section):
    """设置奇数页右侧、偶数页左侧页码。"""
    doc.settings.odd_and_even_pages_header_footer = True
    add_standard_page_number(section.footer, WD_ALIGN_PARAGRAPH.RIGHT)
    add_standard_page_number(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)


def add_ai_disclaimer(doc):
    """在文档最末尾添加 AI 生成提示。"""
    if doc.paragraphs and doc.paragraphs[-1].text.strip() == AI_DISCLAIMER_TEXT:
        return

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(18)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    add_formatted_text(para, AI_DISCLAIMER_TEXT, '仿宋_GB2312', 10.5, color=RGBColor(0x80, 0x80, 0x80))


def set_document_core_properties(doc, title=None):
    """清理并写入基础 DOCX 元数据。"""
    props = doc.core_properties
    now = datetime.now()
    props.title = title or ""
    props.subject = ""
    props.author = ""
    props.keywords = ""
    props.comments = ""
    props.last_modified_by = ""
    props.created = now
    props.modified = now


def get_line_type(line):
    """
    根据Markdown结构判断行类型
    
    返回值:
    - 'title': # 标题（公文标题）
    - 'heading1': ## 一级标题
    - 'heading2': ### 二级标题
    - 'heading3': #### 三级标题
    - 'empty': 空行
    - 'text': 普通文本
    """
    stripped = line.strip()
    
    # 空行
    if not stripped:
        return 'empty'
    
    if stripped.startswith('#### '):
        return 'heading3'
    if stripped.startswith('### '):
        return 'heading2'
    if stripped.startswith('## '):
        return 'heading1'
    if stripped.startswith('# ') and not stripped.startswith('## '):
        return 'title'

    # 兼容模型未加 Markdown 标记但使用公文常见标题格式的情况。
    if re.match(r'^[一二三四五六七八九十]+、[^。；;]{1,40}$', stripped):
        return 'heading1'
    if re.match(r'^（[一二三四五六七八九十]+）[^。；;]{1,40}$', stripped):
        return 'heading2'
    if re.match(r'^\d+[\.．、][^。；;]{1,40}$', stripped):
        return 'heading3'
    
    return 'text'


def strip_markdown_heading(line):
    """去掉Markdown标题标记"""
    stripped = line.strip()
    if stripped.startswith('#### '):
        return stripped[5:]
    if stripped.startswith('### '):
        return stripped[4:]
    if stripped.startswith('## '):
        return stripped[3:]
    if stripped.startswith('# ') and not stripped.startswith('## '):
        return stripped[2:]
    return stripped


def has_markdown_title(lines):
    """判断正文是否已经包含 Markdown 公文标题。"""
    return any(get_line_type(line.strip()) == 'title' for line in lines)


def infer_title_from_output_path(output_path):
    """当正文遗漏 # 标题时，从输出文件名推断公文标题。"""
    if not output_path:
        return None
    stem = Path(os.path.expanduser(output_path)).stem.strip()
    if not stem:
        return None
    stem = re.sub(r'_v\d+$', '', stem)
    stem = re.sub(r'_红头$', '', stem)
    if re.match(r'^公文_\d{8}_\d{6}$', stem):
        return None
    if len(stem) < 4 or len(stem) > 80:
        return None
    return stem


def is_standalone_source_note(output_path):
    """判断当前输出是否为独立的素材来源说明文档。"""
    if not output_path:
        return False
    stem = Path(os.path.expanduser(output_path)).stem.strip()
    return bool(re.search(r'_素材来源说明(?:_v\d+)?$', stem, re.IGNORECASE))


def extract_document_title(lines):
    """从正文行中提取公文标题，用于默认输出文件名。"""
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if get_line_type(stripped) == 'title':
            title = strip_markdown_heading(stripped)
        elif is_plain_document_title(stripped):
            title = stripped
        else:
            continue
        title = sanitize_filename_component(title)
        if title:
            return title
    return None


def strip_markdown_bold(text):
    """去掉Markdown加粗标记 **text** 或 __text__"""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    return text


def normalize_appendix_text(text):
    """清理素材附录中的内部写作痕迹。"""
    text = strip_markdown_bold(text.strip())
    text = re.sub(r'^[→➜]\s*用于[:：]\s*', '用途：', text)
    text = re.sub(r'^[→➜]\s*参考[:：]\s*', '参考：', text)
    text = re.sub(r'^[→➜]\s*用途[:：]\s*', '用途：', text)
    while '来源：来源：' in text:
        text = text.replace('来源：来源：', '来源：')
    while '数据源：数据源：' in text:
        text = text.replace('数据源：数据源：', '数据源：')
    return text


def is_attachment_line(line):
    """判断是否为附件行"""
    return bool(re.match(r'^附件[:：]', line.strip()))


def is_attachment_continuation(line):
    """判断是否为附件清单续行，如 2.×××。"""
    return bool(re.match(r'^\d+[\.．、]', line.strip()))


def set_attachment_list_format(para, continuation=False, line_spacing=None):
    """设置正文末尾附件说明格式。"""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.ind
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:start'), "0")
    ind.set(qn('w:startChars'), "0")
    ind.set(qn('w:end'), "0")
    ind.set(qn('w:endChars'), "0")
    if continuation:
        para.paragraph_format.left_indent = Cm(0)
        para.paragraph_format.first_line_indent = None
        ind.set(qn('w:firstLineChars'), "112")
    else:
        para.paragraph_format.first_line_indent = Cm(0.74)
        ind.set(qn('w:firstLineChars'), "200")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def is_attachment_section_title(line):
    """判断附件正文首页标题。"""
    return line.strip() == "附件"


def document_has_visible_content(doc):
    """判断文档中是否已经有正文内容，用于附件页前分页。"""
    return any(paragraph.text.strip() for paragraph in doc.paragraphs)


def paragraph_has_page_break(paragraph):
    """判断段落中是否包含分页符。"""
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def document_ends_with_page_break(doc):
    """判断当前文档末尾是否已经是分页符，避免重复插入空白页。"""
    for paragraph in reversed(doc.paragraphs):
        if paragraph_has_page_break(paragraph):
            return True
        if paragraph.text.strip():
            return False
    return False


def add_page_break_if_needed(doc, line_spacing=None):
    """在已有正文后插入分页符；若末尾已有分页符则不重复插入。"""
    if not document_has_visible_content(doc) or document_ends_with_page_break(doc):
        return
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)
    if line_spacing:
        set_paragraph_format(para, first_line_indent=False, line_spacing=line_spacing)


def is_signing_entity(line):
    """判断是否为落款单位
    
    落款单位特征：
    - 长度小于30
    - 以机构名结尾（局、厅、委、办、中心等）
    - 不包含具体场所（会议室、办公室等）
    - 不以句号结尾
    """
    stripped = line.strip()
    
    if len(stripped) >= 30:
        return False
    
    # 排除以句号结尾的（不是落款单位）
    if stripped.endswith('。'):
        return False
    
    # 排除包含具体场所的（会议室、服务中心等）
    exclude_keywords = ['会议室', '服务中心', '办事大厅', '窗口']
    if any(kw in stripped for kw in exclude_keywords):
        return False
    
    # 必须以机构名结尾
    agency_suffixes = ['局', '厅', '委', '办', '办公室', '中心', '院', '会', '组委会', '协会', '站', '所', '部', '处', '司', '署', '公司', '集团', '单位']
    if not any(stripped.endswith(suffix) for suffix in agency_suffixes):
        return False
    
    return True


def is_date_line(line):
    """判断是否为公文落款日期。"""
    stripped = re.sub(r'\s+', '', line.strip())
    patterns = [
        r'^\d{4}年(?:\d{1,2}|【[^】]+】)月(?:\d{1,2}|【[^】]+】)日$',
        r'^【[^】]+】年【[^】]+】月【[^】]+】日$',
        r'^(?:\d{4}|×{4}|XXXX)年(?:\d{1,2}|×{1,2}|XX)月(?:\d{1,2}|×{1,2}|XX)日$',
    ]
    return any(re.match(pattern, stripped) for pattern in patterns)


def should_right_align_date(lines, idx):
    """仅将正文落款附近的日期右对齐，避免正文日期误判。"""
    if not is_date_line(lines[idx].strip()):
        return False

    prev_non_empty = ""
    for prev_idx in range(idx - 1, -1, -1):
        candidate = lines[prev_idx].strip()
        if candidate:
            prev_non_empty = candidate
            break

    if prev_non_empty and is_signing_entity(prev_non_empty):
        return True

    next_non_empty = ""
    for next_idx in range(idx + 1, min(len(lines), idx + 4)):
        candidate = lines[next_idx].strip()
        if candidate:
            next_non_empty = candidate
            break

    return next_non_empty in ("[分页符]", "【素材使用情况】", "【知识专库链接】")


def is_contact_info(line):
    """判断是否为联系人信息"""
    return bool(re.match(r'^(?:（?联系人|联系电话|联系人电话|联系方式)[:：]', line.strip()))


def normalize_content_text(content_text):
    """规范输入换行，降低命令行传参导致整篇文本变成一段的风险。"""
    content_text = content_text.replace('\ufeff', '')
    content_text = content_text.replace('\r\n', '\n').replace('\r', '\n')
    content_text = content_text.replace('\u2028', '\n').replace('\u2029', '\n')

    # 如果整段文本被转义成了字面量 \n，恢复为真实换行。
    if '\n' not in content_text and '\\n' in content_text:
        content_text = content_text.replace('\\n', '\n')

    return content_text


def is_plain_document_title(line):
    """判断无 Markdown 标记时首个非空行是否像公文标题。"""
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    if stripped.endswith(('。', '；', ';', '：', ':')):
        return False
    title_suffixes = ('通知', '报告', '请示', '批复', '函', '意见', '方案', '总结', '纪要', '倡议书')
    return stripped.endswith(title_suffixes)


def promote_plain_title(lines):
    """没有 # 标题时，将首个疑似公文标题的非空行提升为标题。"""
    for idx, line in enumerate(lines):
        if not line.strip():
            continue
        if is_plain_document_title(line):
            lines[idx] = f"# {line.strip()}"
        return lines
    return lines


def is_legacy_ai_disclaimer(line):
    """过滤模型自行追加的旧版 AI 提示，统一由脚本在末尾生成规范提示。"""
    stripped = line.strip()
    return stripped in {
        "---",
        "内容由AI生成，仅供参考",
        "内容由 AI 生成，仅供参考",
        AI_DISCLAIMER_TEXT,
    }


def is_recipient(line):
    """判断是否为主送机关
    
    主送机关特征：
    - 包含：政府、厅、局、委、办公室、机构等
    - 以逗号、顿号分隔多个单位
    - 以冒号结尾
    """
    stripped = line.strip()
    
    # 必须以冒号结尾
    if not stripped.endswith('：') and not stripped.endswith(':'):
        return False

    # 排除明显不是主送机关的句子
    # 如："现就...通知如下："、"特此通知："等
    exclude_patterns = [
        r'通知如下：?$',
        r'决定如下：?$',
        r'批复如下：?$',
        r'意见如下：?$',
        r'复函如下：?$',
        r'函复如下：?$',
        r'报告如下：?$',
        r'请示如下：?$',
    ]
    if any(re.search(p, stripped) for p in exclude_patterns):
        return False

    # 占位示例或短主送机关也应识别，如“×××：”
    if len(stripped) <= 30:
        return True

    # 包含主送机关关键词
    keywords = ['政府', '厅', '局', '委', '办公室', '机构', '中心', '公司', '集团']
    has_keyword = any(kw in stripped for kw in keywords)
    
    if not has_keyword:
        return False
    
    # 排除过长的句子（主送机关通常不超过50字）
    if len(stripped) > 50:
        return False
    
    return True


def validate_input(content_text: str) -> bool:
    """验证输入内容"""
    if len(content_text) > MAX_INPUT_LENGTH:
        raise ValueError(f"输入内容过长（{len(content_text):,} 字符）")
    lines = content_text.split('\n')
    if len(lines) > MAX_INPUT_LINES:
        raise ValueError(f"输入行数过多（{len(lines):,} 行）")
    return True


def fix_reference_format(content_text: str) -> str:
    """修复引用格式"""
    def replace_ref(match):
        ref_num = match.group(1)
        return f'[^{ref_num}^]'
    content_text = re.sub(r'\[\^(\d+)\](?!\^)', replace_ref, content_text)
    content_text = re.sub(r'\[\^(\d+)\](?=\[\^\d+\])', replace_ref, content_text)
    return content_text


LANDSCAPE_TABLE_MARKER = "<!-- landscape-table -->"
WIDE_TABLE_COLUMN_THRESHOLD = 6


def parse_markdown_table(lines, start_idx):
    """
    从指定行开始解析Markdown表格，返回 (table_rows, alignments, next_idx)。
    
    支持格式：
    | 列1 | 列2 | 列3 |
    |-----|-----|-----|
    | 数据 | 数据 | 数据 |
    
    Returns:
        list[list[str]]: 解析后的单元格内容（不含分隔行）
        list[str|None]: Markdown 分隔行中提取的列对齐方式
        int: 表格结束后的下一行索引；若start_idx不是表格行则返回 ([], start_idx)
    """
    if start_idx >= len(lines):
        return [], [], start_idx
    
    stripped = lines[start_idx].strip()
    # 表格行必须以 | 开头（可选前导空格）
    if not stripped.startswith('|'):
        return [], [], start_idx
    
    def split_table_row(line):
        """拆分表格行，去掉首尾 | 并分割"""
        s = line.strip()
        if s.startswith('|'):
            s = s[1:]
        if s.endswith('|'):
            s = s[:-1]
        return [cell.strip() for cell in s.split('|')]
    
    def is_separator(line):
        """判断是否为分隔行（如 |---|---|---|）"""
        s = line.strip()
        if not s.startswith('|'):
            return False
        cells = split_table_row(s)
        return all(re.match(r'^[-:]+$', c) for c in cells)

    def parse_alignments(separator_line):
        alignments = []
        for cell in split_table_row(separator_line):
            left = cell.startswith(':')
            right = cell.endswith(':')
            if left and right:
                alignments.append('center')
            elif right:
                alignments.append('right')
            elif left:
                alignments.append('left')
            else:
                alignments.append(None)
        return alignments
    
    rows = []
    alignments = []
    i = start_idx
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped.startswith('|'):
            break
        if is_separator(stripped):
            if not alignments:
                alignments = parse_alignments(stripped)
            i += 1
            continue
        rows.append(split_table_row(stripped))
        i += 1
    
    # 至少要有2行（标题行 + 数据行）才算有效表格
    if len(rows) < 2:
        return [], [], start_idx
    
    return rows, alignments, i


def section_usable_width_dxa(section):
    """获取当前 section 可用正文宽度，单位 DXA/twips。"""
    return int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)


def configure_section(section, landscape=False):
    """按公文版心配置 section，支持横向 A4 宽表页。"""
    page_width_mm, page_height_mm = get_page_size()
    margin_top, margin_bottom, margin_left, margin_right = get_page_margin()
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(page_height_mm)
        section.page_height = Mm(page_width_mm)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(page_width_mm)
        section.page_height = Mm(page_height_mm)
    section.top_margin = Cm(margin_top / 10)
    section.bottom_margin = Cm(margin_bottom / 10)
    section.left_margin = Cm(margin_left / 10)
    section.right_margin = Cm(margin_right / 10)


def add_landscape_section(doc):
    """新增横向 A4 section，并设置国标页码。"""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    add_gbt_page_numbers(doc, section)
    return section


def add_portrait_section(doc):
    """宽表结束后恢复竖向 A4 section。"""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    add_gbt_page_numbers(doc, section)
    return section


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(width_dxa))


def set_table_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn('w:tblCellMar'))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {
        'top': top,
        'start': start,
        'bottom': bottom,
        'end': end,
    }.items():
        node = tbl_cell_mar.find(qn(f'w:{margin_name}'))
        if node is None:
            node = OxmlElement(f'w:{margin_name}')
            tbl_cell_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_grid(table, col_widths):
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement('w:tblGrid')
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in col_widths:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(width))
        tbl_grid.append(grid_col)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def set_row_cant_split(row):
    """尽量禁止 Word 将同一表格行拆到两页。"""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn('w:cantSplit'))
    if cant_split is None:
        cant_split = OxmlElement('w:cantSplit')
        tr_pr.append(cant_split)


def set_table_rows_cant_split(table):
    """尽量保证同一单元格内容不跨页拆分。"""
    for row in table.rows:
        set_row_cant_split(row)


def set_paragraph_keep_with_next(paragraph, keep=True):
    """让表题尽量和后续表格留在同一页。"""
    paragraph.paragraph_format.keep_with_next = keep


def text_width_score(text):
    """估算单元格内容宽度；中文按 2，ASCII 按 1。"""
    score = 0
    for char in str(text):
        score += 1 if ord(char) < 128 else 2
    return score


def compute_column_widths(table_rows, usable_width_dxa):
    num_cols = max(len(row) for row in table_rows)
    scores = []
    for col_idx in range(num_cols):
        values = [row[col_idx] if col_idx < len(row) else '' for row in table_rows]
        max_score = max(text_width_score(value) for value in values) if values else 4
        avg_score = sum(text_width_score(value) for value in values) / max(len(values), 1)
        scores.append(max(6, min(28, max_score * 0.65 + avg_score * 0.35)))

    total_score = sum(scores) or num_cols
    min_width = 900 if num_cols <= 6 else 700
    widths = [max(min_width, int(usable_width_dxa * score / total_score)) for score in scores]
    total_width = sum(widths)
    if total_width != usable_width_dxa and total_width > 0:
        widths[-1] += usable_width_dxa - total_width
    return widths


def infer_column_alignment(table_rows, alignments, col_idx):
    if col_idx < len(alignments) and alignments[col_idx]:
        return alignments[col_idx]

    values = [row[col_idx].strip() for row in table_rows[1:] if col_idx < len(row) and row[col_idx].strip()]
    if values and all(re.match(r'^[+-]?\d+(?:\.\d+)?%?$|^[\d,，.]+$', value) for value in values):
        return 'right'
    if values and max(text_width_score(value) for value in values) <= 12:
        return 'center'
    return 'left'


def alignment_to_word(alignment):
    return {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)


def should_use_landscape_table(table_rows, forced_landscape=False):
    if forced_landscape:
        return True
    num_cols = max(len(row) for row in table_rows)
    return num_cols > WIDE_TABLE_COLUMN_THRESHOLD


def is_table_caption(text):
    """识别表题，如“表1 五省粮食产量对比”。"""
    return bool(re.match(r'^表\s*\d+\s+.+', text.strip()))


def delete_paragraph(paragraph):
    """从文档中删除指定段落。"""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def pop_trailing_table_caption(doc):
    """
    如果文档末尾最近一个可见段落是表题，则移除并返回其文本。
    用于宽表自动横排时，把已写入竖版页的表题移入横版页。
    """
    for paragraph in reversed(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if is_table_caption(text):
            delete_paragraph(paragraph)
            return text
        return None
    return None


def add_table_caption(doc, caption, body_font, body_size):
    """添加表题。表题是普通文字，不设置大纲级别。"""
    if not caption:
        return
    para = doc.add_paragraph()
    add_formatted_text(para, caption, body_font, body_size)
    set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=body_size + 12)
    set_paragraph_keep_with_next(para)


def next_table_number(doc):
    """根据已写入文档的表题计算下一个表格编号。"""
    numbers = []
    for paragraph in doc.paragraphs:
        match = re.match(r'^表\s*(\d+)\s+.+', paragraph.text.strip())
        if match:
            numbers.append(int(match.group(1)))
    return max(numbers, default=0) + 1


def fallback_table_caption(doc):
    """当模型漏写表题时生成兜底表题，避免无名表格进入正式 Word。"""
    return f"表{next_table_number(doc)} 主要情况汇总"


def add_word_table(doc, table_rows, alignments, body_font, body_size, landscape=False, caption=None):
    """
    将解析后的表格数据添加为Word原生表格。
    
    首行作为表头（加粗），其余为数据行。
    所有单元格使用公文字体。
    
    Args:
        doc: Document 对象
        table_rows: list[list[str]], 解析后的表格数据
        body_font: 正文字体名
        body_size: 正文字号(pt)
        landscape: 是否为宽表单独使用横向 A4 页面
        caption: 随表格一起写入的表题
    """
    num_cols = max(len(row) for row in table_rows)
    table_font_size = 15 if landscape or num_cols > WIDE_TABLE_COLUMN_THRESHOLD else body_size

    if landscape:
        caption = caption or pop_trailing_table_caption(doc)
        caption = caption or fallback_table_caption(doc)
        add_landscape_section(doc)
        add_table_caption(doc, caption, body_font, body_size)
    else:
        caption = caption or pop_trailing_table_caption(doc)
        caption = caption or fallback_table_caption(doc)
        add_table_caption(doc, caption, body_font, body_size)
    
    table = doc.add_table(rows=len(table_rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    usable_width = section_usable_width_dxa(doc.sections[-1])
    col_widths = compute_column_widths(table_rows, usable_width)
    set_table_width(table, usable_width)
    set_table_grid(table, col_widths)
    set_table_cell_margins(table)
    
    for row_idx, row_data in enumerate(table_rows):
        for col_idx in range(num_cols):
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, col_widths[col_idx])
            # 清空默认段落
            cell.paragraphs[0].clear()
            text = row_data[col_idx] if col_idx < len(row_data) else ''
            alignment = 'center' if row_idx == 0 else infer_column_alignment(table_rows, alignments, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = alignment_to_word(alignment)
            add_formatted_text(paragraph, text, body_font, table_font_size, bold=(row_idx == 0))
            set_paragraph_format(
                paragraph,
                first_line_indent=False,
                alignment=alignment_to_word(alignment),
                line_spacing=table_font_size + 6,
            )
    set_table_rows_cant_split(table)
    
    # 表格后加一个空段落，与后续内容保持间距
    spacer = doc.add_paragraph()
    set_paragraph_format(spacer, first_line_indent=False, line_spacing=body_size + 8)

    if landscape:
        add_portrait_section(doc)


def set_paragraph_mark_font(para, font_name, font_size):
    """设置空段落的段落标记字体，避免空行高度受默认样式影响。"""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    latin_font = get_latin_font_config()
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

    for tag in ('w:sz', 'w:szCs'):
        elem = rPr.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            rPr.append(elem)
        elem.set(qn('w:val'), str(int(font_size * 2)))


def add_blank_paragraphs(doc, count, line_spacing, font_name=None, font_size=None):
    """插入指定数量的空段，用于公文固定空行。"""
    for _ in range(count):
        para = doc.add_paragraph()
        set_paragraph_format(para, first_line_indent=False, line_spacing=line_spacing)
        if font_name and font_size:
            set_paragraph_mark_font(para, font_name, font_size)


def create_document(content_text, output_path=None):
    """
    创建排版后的普通公文文档（无红头）
    
    Args:
        content_text: 正文内容（Markdown格式）
        output_path: 输出文件路径
    
    Returns:
        输出文件路径
    """
    content_text = normalize_content_text(content_text)
    validate_input(content_text)
    content_text = fix_reference_format(content_text)
    
    doc = Document()
    
    # 设置 A4 纸张和页边距（公文格式规范）
    for section in doc.sections:
        configure_section(section, landscape=False)
        add_gbt_page_numbers(doc, section)
    
    # 处理正文内容
    lines = content_text.strip().split('\n')
    default_title = extract_document_title(lines)
    if not has_markdown_title(lines):
        inferred_title = infer_title_from_output_path(output_path)
        before = list(lines)
        lines = promote_plain_title(lines)
        if before == lines and inferred_title:
            lines = [f"# {inferred_title}", ""] + lines
    i = 0
    line_count = len(lines)
    
    # 从配置获取字体信息
    title_font, title_size, _ = get_font_config('title')
    title_line_spacing = FORMAT_CONFIG['title'].get('line_spacing_pt', 33) if FORMAT_CONFIG and 'title' in FORMAT_CONFIG else 33
    body_font, body_size, _ = get_font_config('body')
    h1_font, h1_size, _ = get_font_config('heading1')
    h2_font, h2_size, _ = get_font_config('heading2')
    h3_font, h3_size, h3_bold = get_font_config('heading3')
    body_line_spacing = FORMAT_CONFIG['body']['line_spacing_pt'] if FORMAT_CONFIG else 28
    
    # 独立素材说明不是正文附录，保留专用排版但不在首个栏目之前分页。
    standalone_source_note = is_standalone_source_note(output_path)
    # 附录区域标志：素材使用情况、参考来源和知识专库链接统一靠左排版
    in_appendix_section = False
    # 知识专库链接区域标志：URL 生成可点击蓝链
    in_kb_section = False
    # 知识专库链接前一行标签，用于显示为“标题（点击打开）”
    pending_kb_label = None
    kb_link_count = 0
    # 附件正文首页“附件”后的下一行作为附件标题处理
    awaiting_attachment_title = False
    # 正文末尾附件目录后，落款前需要固定空三行
    attachment_list_pending_sign_gap = False
    # 表格前一行使用 <!-- landscape-table --> 时，强制下一张表横排。
    force_next_table_landscape = False
    pending_landscape_table_caption = None

    while i < line_count:
        line = lines[i]
        stripped = line.strip()
        line_type = get_line_type(stripped)

        if is_legacy_ai_disclaimer(stripped):
            i += 1
            continue

        if stripped == LANDSCAPE_TABLE_MARKER:
            force_next_table_landscape = True
            i += 1
            continue

        if force_next_table_landscape and is_table_caption(stripped):
            pending_landscape_table_caption = stripped
            i += 1
            continue

        # Markdown表格检测（必须在其他判断之前）
        if stripped.startswith('|'):
            table_rows, alignments, next_i = parse_markdown_table(lines, i)
            if table_rows:
                use_landscape = should_use_landscape_table(table_rows, force_next_table_landscape)
                add_word_table(
                    doc,
                    table_rows,
                    alignments,
                    body_font,
                    body_size,
                    landscape=use_landscape,
                    caption=pending_landscape_table_caption if use_landscape else None,
                )
                force_next_table_landscape = False
                pending_landscape_table_caption = None
                i = next_i
                continue
            # 不是有效表格，当普通文本处理，继续往下走

        # 空行
        if line_type == 'empty':
            i += 1
            continue

        if (
            attachment_list_pending_sign_gap
            and not is_attachment_line(stripped)
            and not is_attachment_continuation(stripped)
        ):
            if is_signing_entity(stripped):
                add_blank_paragraphs(doc, 3, body_line_spacing, body_font, body_size)
            elif is_date_line(stripped):
                add_blank_paragraphs(doc, 3, body_line_spacing, body_font, body_size)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_formatted_text(para, stripped, body_font, body_size)
                set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=body_line_spacing)
                attachment_list_pending_sign_gap = False
                i += 1
                continue
            attachment_list_pending_sign_gap = False

        if awaiting_attachment_title:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(para, stripped, title_font, title_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=title_line_spacing)
            set_outline_level(para, 0)
            awaiting_attachment_title = False
            i += 1
            continue

        marker_text = strip_markdown_heading(stripped)
        if marker_text in ("素材使用情况", "参考资料", "参考来源"):
            if not standalone_source_note:
                add_page_break_if_needed(doc, body_line_spacing)
            in_appendix_section = True
            para = doc.add_paragraph()
            label = "【素材使用情况】" if marker_text == "素材使用情况" else f"【{marker_text}】"
            add_formatted_text(para, label, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            i += 1
            continue

        if marker_text == "知识专库链接":
            if not in_appendix_section and not standalone_source_note:
                add_page_break_if_needed(doc, body_line_spacing)
            in_appendix_section = True
            in_kb_section = True
            pending_kb_label = None
            kb_link_count = 0
            para = doc.add_paragraph()
            add_formatted_text(para, "【知识专库链接】", body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 检测附录区域开始
        if (
            re.match(r'^【素材使用情况】', stripped)
            or re.match(r'^【参考资料】', stripped)
            or re.match(r'^【参考来源】', stripped)
            or stripped in ("素材使用情况", "参考资料", "参考来源")
        ):
            if not in_appendix_section and not standalone_source_note:
                add_page_break_if_needed(doc, body_line_spacing)
            in_appendix_section = True

        # 检测【知识专库链接】区域开始
        if re.match(r'^【知识专库链接】', stripped):
            if not in_appendix_section and not standalone_source_note:
                add_page_break_if_needed(doc, body_line_spacing)
            in_appendix_section = True
            in_kb_section = True
            pending_kb_label = None
            kb_link_count = 0

        # 附录区域：统一靠左排版；知识专库 URL 生成可点击蓝链。
        # 必须早于标题识别，避免“15. XXX”这类素材编号被误判为三级标题。
        if in_appendix_section:
            if in_kb_section and re.search(r'https?://', stripped):
                url_match = re.search(r'(https?://\S+)', stripped)
                if url_match:
                    url = sanitize_url(url_match.group(1))
                    if is_internal_search_endpoint(url):
                        i += 1
                        continue
                    kb_link_count += 1
                    para = doc.add_paragraph()
                    add_hyperlink(
                        para,
                        url,
                        hyperlink_display_text(pending_kb_label, kb_link_count),
                        font_name=body_font,
                        font_size=body_size,
                    )
                    set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
                    pending_kb_label = None
                    i += 1
                    continue

            para = doc.add_paragraph()
            clean_text = normalize_appendix_text(stripped)
            if in_kb_section and is_useless_kb_label_line(clean_text):
                i += 1
                continue
            add_formatted_text(para, clean_text, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            if in_kb_section and clean_text and not clean_text.startswith("【") and not is_generic_kb_label(clean_text):
                pending_kb_label = clean_text
            i += 1
            continue

        # 附件清单续行要先于三级标题识别，避免 “2.附件名” 被误判为三级标题。
        # 仅在已识别到正文末尾“附件：”清单后生效，避免普通正文连续编号被误排为附件续行。
        if attachment_list_pending_sign_gap and i > 0 and is_attachment_continuation(stripped):
            prev_non_empty = ""
            for prev_idx in range(i - 1, -1, -1):
                prev_non_empty = lines[prev_idx].strip()
                if prev_non_empty:
                    break
            if is_attachment_line(prev_non_empty) or is_attachment_continuation(prev_non_empty):
                para = doc.add_paragraph()
                add_formatted_text(para, f"        {stripped}", body_font, body_size)
                set_attachment_list_format(para, continuation=True, line_spacing=body_line_spacing)
                attachment_list_pending_sign_gap = True
                i += 1
                continue

        # 公文标题（# 标题）
        if line_type == 'title':
            content = strip_markdown_heading(stripped)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_formatted_text(para, content, title_font, title_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=title_line_spacing)
            set_outline_level(para, 0)  # 文档标题设为1级
            add_blank_paragraphs(doc, 1, body_line_spacing, body_font, body_size)
            i += 1
            continue

        # 一级标题（## 一、XXX）
        if line_type == 'heading1':
            content = strip_markdown_heading(stripped)
            para = doc.add_paragraph()
            add_formatted_text(para, content, h1_font, h1_size)
            set_paragraph_format(para, indent_chars=2, line_spacing=body_line_spacing)
            set_outline_level(para, 1)  # 一级标题设为2级
            i += 1
            continue

        # 二级标题（### （一）XXX）
        if line_type == 'heading2':
            content = strip_markdown_heading(stripped)
            para = doc.add_paragraph()
            add_formatted_text(para, content, h2_font, h2_size)
            set_paragraph_format(para, indent_chars=2, line_spacing=body_line_spacing)
            set_outline_level(para, 2)  # 二级标题设为3级
            i += 1
            continue

        # 三级标题（#### 1.XXX）
        if line_type == 'heading3':
            content = strip_markdown_heading(stripped)
            para = doc.add_paragraph()
            add_formatted_text(para, content, h3_font, h3_size, bold=h3_bold)
            set_paragraph_format(para, indent_chars=2, line_spacing=body_line_spacing)
            set_outline_level(para, 3)  # 三级标题设为4级
            i += 1
            continue

        # 附件
        if is_attachment_line(stripped):
            add_blank_paragraphs(doc, 1, body_line_spacing, body_font, body_size)
            para = doc.add_paragraph()
            add_formatted_text(para, stripped, body_font, body_size)
            set_attachment_list_format(para, continuation=False, line_spacing=body_line_spacing)
            attachment_list_pending_sign_gap = True
            i += 1
            continue

        if is_attachment_section_title(stripped):
            if document_has_visible_content(doc):
                page_break_para = doc.add_paragraph()
                page_break_run = page_break_para.add_run()
                page_break_run.add_break(WD_BREAK.PAGE)
                set_paragraph_format(page_break_para, first_line_indent=False, line_spacing=body_line_spacing)
            para = doc.add_paragraph()
            add_formatted_text(para, stripped, h1_font, h1_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            awaiting_attachment_title = True
            i += 1
            continue

        # 落款单位
        if is_signing_entity(stripped):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_formatted_text(para, stripped, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 联系信息
        if is_contact_info(stripped):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_formatted_text(para, stripped, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 落款日期：优先根据前一行落款单位和后续附录标记判断，避免素材页影响位置比例
        if should_right_align_date(lines, i):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_formatted_text(para, stripped, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 分页符标记
        if stripped == '[分页符]':
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_break(WD_BREAK.PAGE)
            i += 1
            continue

        # 主送机关（前10行内）
        if is_recipient(stripped) and i < 10:
            para = doc.add_paragraph()
            add_formatted_text(para, stripped, body_font, body_size)
            set_paragraph_format(para, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=body_line_spacing)
            i += 1
            continue

        # 普通正文
        # 清理Markdown加粗标记
        clean_text = strip_markdown_bold(stripped)
        para = doc.add_paragraph()
        add_formatted_text(para, clean_text, body_font, body_size)
        set_paragraph_format(para, indent_chars=2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=body_line_spacing)
        i += 1
    
    # 保存文件
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    if not output_path:
        default_stem = default_title or f'公文_{timestamp}'
        output_path = os.path.join(get_configured_output_dir(), f'{default_stem}.docx')

    output_path = resolve_output_path(output_path)

    if not output_path.lower().endswith('.docx'):
        output_path = output_path.rsplit('.', 1)[0] + '.docx' if '.' in output_path else output_path + '.docx'

    # 添加版本号（如果文件已存在）
    output_path = get_next_version(output_path)

    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    add_ai_disclaimer(doc)
    set_document_core_properties(doc, default_title)
    doc.save(output_path)
    return display_path(output_path)


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='公文排版工具 v2.0.0 - 只支持普通格式')
    parser.add_argument('input', nargs='?', help='输入文件路径')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--text', help='直接输入公文文本（如果是文件路径会自动读取）')
    
    args = parser.parse_args()
    
    if args.text:
        # 智能检测：如果 --text 是文件路径，自动读取文件内容
        try:
            text_path = resolve_input_text_path(args.text)
        except ValueError:
            text_path = None
        if text_path and os.path.exists(text_path) and os.path.isfile(text_path):
            print(f'⚠ 检测到 --text 参数是文件路径，自动读取文件内容: {text_path}')
            with open(text_path, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            content = args.text
    elif args.input:
        input_path = resolve_input_text_path(args.input)
        if os.path.exists(input_path):
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            raise FileNotFoundError(f"输入文件不存在: {input_path}")
    else:
        parser.print_help()
        sys.exit(1)
    
    try:
        output_path = create_document(
            content,
            output_path=args.output
        )
        
        print(f'✓ 普通公文已生成: {output_path}')
    except Exception as e:
        print(f'✗ 生成失败: {e}')
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    main()
